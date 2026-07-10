# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt

path = r"c:\Users\LENOVO\Desktop\实验一 iPhone X写实三维建模实验报告.docx"
doc = Document(path)

print("=== SECTIONS ===")
for i, s in enumerate(doc.sections):
    print(f"Section {i}: margins L={s.left_margin} R={s.right_margin} T={s.top_margin} B={s.bottom_margin}")

print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "None"
    text = p.text.strip()
    if not text and style == "Normal":
        continue
    print(f"{i:3d} | style={style:25s} | align={p.alignment} | {text[:100]}")
    for r in p.runs[:2]:
        sz = r.font.size
        print(f"      run: name={r.font.name} size={sz} bold={r.bold} text={r.text[:40]!r}")

print(f"\n=== TABLES: {len(doc.tables)} ===")
for ti, t in enumerate(doc.tables):
    print(f"\nTable {ti}: {len(t.rows)}x{len(t.columns)}")
    for ri, row in enumerate(t.rows[:12]):
        cells = [c.text.strip().replace("\n", " ")[:40] for c in row.cells]
        print(f"  R{ri}: {cells}")
