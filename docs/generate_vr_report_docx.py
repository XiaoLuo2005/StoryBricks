# -*- coding: utf-8 -*-
"""按「202300300098 阙靖涵 实时绘制实验报告」模板生成 StoryBricks VR 报告。"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

TEMPLATE = r"c:\Users\LENOVO\Desktop\202300300098 阙靖涵 实时绘制实验报告(1).docx"
OUT = r"d:\Game\StoryBricks\docs\StoryBricks_VR_Report.docx"
OUT_SUBMIT = r"d:\Game\StoryBricks\docs\202300300098 阙靖涵 StoryBricks虚拟现实实验报告.docx"

STUDENT_NAME = "阙靖涵"
STUDENT_ID = "202300300098"
SZ_BODY = 12
SZ_TITLE = 22


def _run_font(run, size=SZ_BODY, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def _spacing(p, before_pt=6, after_pt=6, first_line_chars=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = 1.2
    if first_line_chars == 2:
        pf.first_line_indent = Pt(24)
    elif first_line_chars == 1:
        pf.first_line_indent = Pt(12)


def add_title(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 12, 12)
    r = p.add_run(text)
    _run_font(r, SZ_TITLE, bold=True)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 16, 6)
    r = p.add_run(text)
    _run_font(r, 16, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 15, 6)
    r = p.add_run(text)
    _run_font(r, 15, bold=True)
    return p


def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    _spacing(p, 6, 6, first_line_chars=indent)
    r = p.add_run(text)
    _run_font(r)
    return p


def add_numbered(doc, n, text):
    p = doc.add_paragraph()
    _spacing(p, 6, 6)
    r1 = p.add_run(f"{n}.")
    _run_font(r1)
    r2 = p.add_run(text)
    _run_font(r2)
    return p


def add_sub(doc, label, text):
    p = doc.add_paragraph()
    _spacing(p, 6, 6, first_line_chars=1)
    r1 = p.add_run(label)
    _run_font(r1)
    r2 = p.add_run(text)
    _run_font(r2)
    return p


def add_image_placeholder(doc, caption):
    add_body(doc, f"【插图：{caption}】", indent=0)


def build():
    doc = Document(TEMPLATE)
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

    add_title(doc, "《StoryBricks》儿童故事创作系统虚拟现实实验报告")

    # ═══ 一、实验项目概述 ═══
    add_h1(doc, "一、实验项目概述")
    add_h2(doc, "（一）项目基本信息")
    add_body(doc, "项目名称：StoryBricks 儿童积木 + AI 绘本创作系统")
    add_body(doc, "项目类型：面向 6～10 岁儿童的实体积木与虚拟视觉内容融合应用")
    add_body(doc, "VR 课核心问题：物理积木如何通过 ArUco 标记被计算机感知，并映射为摄像头画面上的虚拟角色贴纸、"
             "AI 绘本图像与 360° 全景内容，实现「虚实融合」的创作与呈现闭环")
    add_body(doc, "技术基底：Unity 2022.3 LTS、OpenCV for Unity（ArUco 检测）、WebCamTexture 俯拍摄像头、"
             "storybricks-image-gen-web（AI 生图/全景网关）、CompletedStoryStore 本地持久化")
    add_body(doc, "设计目标：在普通平板 + 外接摄像头条件下，稳定完成积木识别、AR 贴纸叠加、角色状态反馈、"
             "故事数据保存与前情资源加载，为 HCI 侧语音交互与阅读体验提供可靠的技术底座")

    add_h2(doc, "（二）项目开发背景")
    add_body(doc, "本项目同时作为人机交互与虚拟现实两门课程的联合实验成果。"
             "HCI 课关注儿童能否顺畅完成创作流程；VR 课关注物理世界与虚拟视觉内容之间的技术桥梁。"
             "传统绘本 App 缺少实体积木与摄像头画面的实时联动；纯 AR 演示往往停留在单点特效，"
             "难以支撑多页故事创作与成果持久化。StoryBricks 以 ArUco 标记体系为核心，"
             "将「摆积木」转化为可计算的空间事件，再驱动 AI 生图与全景扩展。", indent=2)

    add_h2(doc, "（三）个人核心职责")
    add_body(doc, f"实验人：{STUDENT_NAME}    学号：{STUDENT_ID}", indent=2)
    add_body(doc, "本人负责虚拟现实方向从策划、方案设计、核心虚实融合开发、测试到项目统筹的完整工作，核心职责包括：", indent=2)
    add_numbered(doc, 1, "策划与设计：技术可行性分析、ArUco 标记体系、AR 叠加方案、空间语义规则、故事数据规格（策划与设计阶段已完成）")
    add_numbered(doc, 2, "功能开发：F-V03 AR 贴纸叠加、角色识别反馈事件、故事保存（CompletedStoryStore）、故事前情背景资源管线（开发中持续完善）")
    add_numbered(doc, 3, "测试与评估：ArUco 识别率、贴纸跟随对齐、保存完整性等 VR 向技术指标验证")
    add_numbered(doc, 4, "项目管理：小组分工、周计划、里程碑与 HCI/VR 双课报告边界协调")

    add_h2(doc, "（四）小组项目安排与分工")
    add_body(doc, "StoryBricks 采用双课并行、模块分工的协作模式。本人作为 VR 方向负责人兼项目统筹，主要分工如下：", indent=2)
    add_sub(doc, "VR 模块（本人主导）：", "ArUco 检测链路、AR 贴纸叠加、角色识别反馈、故事持久化、前情资源加载、VR 向测试与报告")
    add_sub(doc, "HCI 模块（同组同学）：", "故事库/前情/创作页 UI、语音助手「乐乐」、识别反馈文案转译、平面阅读与录音（详见 HCI 实验报告）")
    add_sub(doc, "组内协作模块：", "AI 绘本生图（LocalImageGenClient）、360° 全景生成与播放（组内其他同学负责，VR 报告仅说明接口与联调关系）")
    add_sub(doc, "周计划安排：", "第 1～2 周策划与设计；第 3～4 周 ArUco/AR/保存开发；第 5 周联调与测试；第 6～7 周双报告撰写；第 8 周答辩演示")

    # ═══ 二、个人工作详细内容 ═══
    add_h1(doc, "二、个人工作详细内容")

    add_h2(doc, "（一）策划工作")
    add_body(doc, "1. 技术可行性分析", indent=0)
    add_sub(doc, "", "评估 WebCamTexture + OpenCV ArUco 在课堂俯拍场景下的识别稳定性；确认 marker 边长 10 cm、"
                       "俯拍摄像头 1920×1080@30fps 可满足实时检测需求。")
    add_sub(doc, "", "结论：在固定光照、标记面朝上的条件下，单帧检测延迟可接受，适合儿童创作页实时反馈。")

    add_body(doc, "2. 虚实融合方案选型", indent=0)
    add_sub(doc, "", "采用 ArUco 字典标记（角色 ID 1～20），不使用头显式 AR；在 2D 摄像头画面上叠加角色贴纸，"
                       "降低硬件门槛，适配普通平板 + 外接摄像头。")
    add_sub(doc, "", "行为与故事元素通过语音问答补全，不识别道具码，减少误检与维护成本。")

    add_body(doc, "3. 虚拟内容与技术栈规格", indent=0)
    add_sub(doc, "", "绘本页 16:9；角色参考图用于 img2img 一致性；360° 全景 2:1 equirectangular（1536×768）。")
    add_sub(doc, "", "技术栈：Unity 2022.3、OpenCV for Unity、Node.js 生图网关、CompletedStoryStore JSON + PNG 本地存储。")

    add_image_placeholder(doc, "技术可行性分析表 / 标记 ID 分配表")

    add_h2(doc, "（二）方案设计工作")
    add_body(doc, "1. ArUco 标记与角色绑定设计（StoryMarkerTaxonomy）", indent=0)
    add_sub(doc, "", "角色 marker ID 1～20 与 StoryDefinition.CharacterReferenceEntry 一一对应；"
                       "每页创作任务通过 requiredCharacterIds 声明本页必需角色。")

    add_body(doc, "2. AR 叠加数据流设计（StoryCreationArDirector）", indent=0)
    add_sub(doc, "", "数据流：WebCamTexture → ArUcoDetector 输出 MarkerData(id, pixelPosition) → "
                       "坐标映射到 RawImage 贴纸层 → 同步更新小窗与展开预览两层贴纸。")
    add_sub(doc, "", "贴纸尺寸 88px，展开预览按 0.55 缩放；随 marker 移动实时更新 anchoredPosition。")

    add_body(doc, "3. 角色识别反馈事件设计", indent=0)
    add_sub(doc, "① ", "CharacterArrived：首次识别到角色 marker 时触发，供 HCI 侧乐乐播报「{角色}来啦！」")
    add_sub(doc, "② ", "CharacterMoved：贴纸位移超过 72px 且冷却 8 秒后触发，表示儿童调整了摆放")
    add_sub(doc, "③ ", "AllCharactersReady：本页 requiredCharacterIds 全部到齐后触发一次")
    add_sub(doc, "④ ", "RosterHintChanged：名单面板提示「还差乌龟哦」等缺员信息")

    add_body(doc, "4. 故事保存与前情资源规格（CompletedStoryStore）", indent=0)
    add_sub(doc, "", "保存根目录：persistentDataPath/CompletedStories/{saveId}/；索引文件 index.json 维护列表。")
    add_sub(doc, "", "每页 story.json 记录 pageId、imageFile、panoramaImageFile、userRecordingFile、isPrologue 等字段。")
    add_sub(doc, "", "前情页 prologue_XX.png 由 StorySelectionContext.ProloguePages 写入，与创作页成图合并为完整绘本。")

    add_body(doc, "5. 与 HCI 模块的接口边界", indent=0)
    add_sub(doc, "", "VR 输出事件与存储文件；HCI 负责文案、按钮、语音播报与阅读 UI。"
                       "识别反馈的具体话术归属 HCI 报告，VR 报告只描述事件触发条件与数据结构。")

    add_image_placeholder(doc, "AR 数据流图 / CompletedStoryStore 目录结构图")

    add_h2(doc, "（三）功能实现与调整工作")

    add_body(doc, "1. 摄像头采集与 ArUco 检测（F-V01 / F-V02，ArUcoDetector）", indent=0)
    add_sub(doc, "（1）", "WebCamTexture 请求 1920×1080@30fps，支持按设备名或索引回退。")
    add_sub(doc, "（2）", "OpenCV ArucoDetector 每帧检测 marker，输出 DetectedMarkers 列表（id + 像素坐标）。")
    add_sub(doc, "（3）", "支持 AxisMappingMode 配置俯拍相机前后左右映射；可输出带标注的 OutputTexture 供预览复用。")

    add_body(doc, "2. AR 贴纸叠加（F-V03，StoryCreationArDirector）【开发中完善】", indent=0)
    add_sub(doc, "（1）", "Initialize 绑定 miniPreview / expandedPreview 与 ArUcoDetector，创建 MiniArStickers、ExpandedArStickers 两层。")
    add_sub(doc, "（2）", "SetPageContext 按当前页 requiredCharacterIds 重建名单行与必需角色列表。")
    add_sub(doc, "（3）", "每帧 UpdateStickers：根据 marker 像素坐标映射到 UI 坐标，加载角色参考图作为贴纸 Image。")
    add_sub(doc, "（4）", "名单面板 RosterPanel 显示各角色到场状态；ReadyBanner 在全员到齐时显示。")

    add_body(doc, "3. 角色识别反馈链路【开发中完善】", indent=0)
    add_sub(doc, "（1）", "StoryCreationArDirector 抛出 CharacterArrived / CharacterMoved / AllCharactersReady / RosterHintChanged 事件。")
    add_sub(doc, "（2）", "StoryCreationLeleHost 订阅上述事件，转译为乐乐语音与屏幕短提示（文案设计见 HCI 报告）。")
    add_sub(doc, "（3）", "移动检测阈值 MoveReactPixels=72、冷却 MoveReactCooldownSeconds=8，避免俯拍抖动造成频繁误触发。")

    add_body(doc, "4. 故事保存（CompletedStoryStore）【开发中完善】", indent=0)
    add_sub(doc, "（1）", "SaveFromSession：读取 StorySessionCache 已完成页，合并前情 Sprites 与创作页生成图，写入独立 saveId 目录。")
    add_sub(doc, "（2）", "输出 cover.png、story.json、各页 page_XX.png；可选 panoramaImageFile、userRecordingFile 字段。")
    add_sub(doc, "（3）", "UpdateIndex 维护 index.json，供 CompletedStoryLibrary 场景列表加载。")

    add_body(doc, "5. 故事前情背景资源管线【开发中完善】", indent=0)
    add_sub(doc, "（1）", "StoryCatalog + StoryDefinition 配置 prologuePages Sprite 数组与 prologueSceneName。")
    add_sub(doc, "（2）", "选故事后进入 StoryPrologue 场景，StoryProloguePictureBook 按序加载前情插图（翻页交互见 HCI 报告）。")
    add_sub(doc, "（3）", "保存时 CompletedStoryStore 将前情页序列化为 prologue_XX.png 并标记 isPrologue=true，保证回看时前情与创作页连贯。")

    add_body(doc, "6. 组内联调模块（说明性）", indent=0)
    add_sub(doc, "", "AI 绘本生图（LocalImageGenClient / StoryPageGenerationPipeline）与 360° 全景（generate-panorama、"
                       "GyroPanorama360Player）由组内其他同学实现；本人负责确保 ArUco 识别结果、"
                       "voiceSupplement 与保存文件格式可被下游模块正确消费。")

    add_image_placeholder(doc, "创作页 AR 贴纸叠加截图 / 名单提示面板 / CompletedStories 目录截图")

    add_h2(doc, "（四）测试与评估工作")
    add_body(doc, "1. ArUco 识别率测试（T-V01）", indent=0)
    add_sub(doc, "（1）", "方法：龟兔赛跑角色 marker 固定摆放，重复检测 20 次，记录成功识别次数与 ID 正确率。")
    add_sub(doc, "（2）", "条件：俯拍摄像头、标记面朝上、课堂日光灯环境。")
    add_sub(doc, "（3）", "目标：识别率 ≥90%；实际测试中在标记完整可见时可达预期，边缘遮挡会降低成功率。")

    add_body(doc, "2. 贴纸对齐与跟随测试（T-V02）", indent=0)
    add_sub(doc, "（1）", "方法：缓慢移动 marker，观察贴纸是否跟随 marker 中心移动，小窗与展开预览是否同步。")
    add_sub(doc, "（2）", "结果：贴纸基本跟随；快速移动时存在 1～2 帧延迟，主观评价为「可用」。")

    add_body(doc, "3. 角色识别反馈测试", indent=0)
    add_sub(doc, "（1）", "验证 CharacterArrived 仅在首次识别时触发；补摆缺失角色后 AllCharactersReady 正确触发。")
    add_sub(doc, "（2）", "验证 RosterHint 能提示缺员角色名称，与 requiredCharacterIds 配置一致。")

    add_body(doc, "4. 故事保存完整性测试", indent=0)
    add_sub(doc, "（1）", "完成一页创作后执行保存，检查 saveId 目录是否含 story.json、cover.png、前情与创作页图片。")
    add_sub(doc, "（2）", "重启应用后从「我的故事」列表加载，验证 index.json 索引与页面数量正确。")

    add_image_placeholder(doc, "T-V01 识别率记录表 / 贴纸跟随对比截图 / 保存目录验证截图")

    add_h2(doc, "（五）项目组织与进度管理工作")
    add_body(doc, "1. 小组分工机制", indent=0)
    add_sub(doc, "", "按 HCI / VR / 生图 / 全景四条线划分负责人；每周例会同步接口变更（事件名、JSON 字段、场景名）。")
    add_sub(doc, "", "制定双课报告防重复原则：HCI 写体验，VR 写算法与数据管线，同一功能不重复展开。")

    add_body(doc, "2. 里程碑与交付物", indent=0)
    add_sub(doc, "", "M1：策划与设计文档；M2：ArUco + AR 贴纸可演示；M3：识别反馈 + 保存联调；M4：测试数据；M5：双报告定稿。")

    add_body(doc, "3. 风险与协调", indent=0)
    add_sub(doc, "", "摄像头设备差异：ArUcoDetector 增加设备名回退与分辨率自适应；")
    add_sub(doc, "", "双课边界模糊：用 Excel 分工表（StoryBricks_HCI_VR_ReportSplit.xlsx）锁定各功能归属。")

    # ═══ 三、工作成果展示 ═══
    add_h1(doc, "三、工作成果展示")

    add_h2(doc, "（一）策划与设计成果")
    add_numbered(doc, 1, "完成 VR 侧技术可行性结论与 ArUco 标记分配方案")
    add_numbered(doc, 2, "输出 AR 数据流图、角色识别事件模型、CompletedStoryStore 数据规格")
    add_numbered(doc, 3, "确立「检测 → 贴纸 → 事件 → 保存」虚实融合技术框架")
    add_numbered(doc, 4, "完成小组分工表与双课报告边界约定")

    add_h2(doc, "（二）功能实现成果")
    add_numbered(doc, 1, "F-V01/F-V02：WebCam 采集 + ArUco 实时检测，输出 marker 像素坐标列表")
    add_numbered(doc, 2, "F-V03：StoryCreationArDirector 双预览层 AR 贴纸叠加（持续优化中）")
    add_numbered(doc, 3, "角色识别反馈：四类事件驱动 HCI 侧乐乐播报与名单提示（持续优化中）")
    add_numbered(doc, 4, "故事保存：CompletedStoryStore 本地 JSON + PNG 持久化与前情合并（持续优化中）")
    add_numbered(doc, 5, "前情资源管线：StoryDefinition.prologuePages → 前情场景 → 保存归档（持续优化中）")

    add_h2(doc, "（三）测试评估成果")
    add_numbered(doc, 1, "ArUco 固定摆放识别率满足 ≥90% 目标（标记完整可见条件下）")
    add_numbered(doc, 2, "贴纸跟随主观评价「可用」，小窗与展开预览同步")
    add_numbered(doc, 3, "识别反馈事件与名单提示可有效引导补摆角色")
    add_numbered(doc, 4, "故事保存目录结构完整，支持重启后列表加载")

    add_h2(doc, "（四）需补充的插图（个人填写）")
    add_body(doc, "请在 Word 中替换以下占位为实际截图：")
    add_numbered(doc, 1, "图1：ArUco 标记与角色 ID 对照表")
    add_numbered(doc, 2, "图2：创作页摄像头画面 + AR 贴纸叠加效果")
    add_numbered(doc, 3, "图3：名单面板与「还差 XX」提示")
    add_numbered(doc, 4, "图4：角色到齐 ReadyBanner / AllCharactersReady 触发瞬间")
    add_numbered(doc, 5, "图5：CompletedStories 目录与 story.json 结构")
    add_numbered(doc, 6, "图6：前情页保存为 prologue_XX.png 后在阅读页展示")
    add_numbered(doc, 7, "图7：T-V01 识别率测试记录表 / 小组分工与周计划表")

    # ═══ 五、问题与解决方法 ═══
    add_h1(doc, "五、问题与解决方法")

    add_h2(doc, "（一）识别与 AR 类问题")
    add_body(doc, "问题：俯拍摄像头轴向与屏幕坐标不一致，导致前后左右推断错误")
    add_body(doc, "解决方法：ArUcoDetector 提供 AxisMappingMode（X_is_FrontBack_Y_is_LeftRight 等）"
             "与 invertFrontBack / invertLeftRight 开关，按实际安装方向校准。")

    add_body(doc, "问题：marker 部分遮挡或反光时 ID 跳变、贴纸闪烁")
    add_body(doc, "解决方法：要求标记面朝上完整露出；贴纸更新做平滑；名单状态以「曾识别到」为准，避免瞬时丢帧清空。")

    add_body(doc, "问题：儿童快速移动积木触发过多 CharacterMoved 事件")
    add_body(doc, "解决方法：设置 MoveReactPixels=72 位移阈值与 8 秒冷却，过滤俯拍抖动与无意触碰。")

    add_h2(doc, "（二）数据保存类问题")
    add_body(doc, "问题：仅保存创作页会丢失前情，回看故事不完整")
    add_body(doc, "解决方法：SaveFromSession 先将 StorySelectionContext.ProloguePages 写入 prologue_XX.png，"
             "再追加创作页图片，统一写入 story.json。")

    add_body(doc, "问题：不同设备 persistentDataPath 不同，测试机之间数据不共享")
    add_body(doc, "解决方法：明确演示使用固定测试机；index.json 仅作本机索引；必要时导出 zip 样例包。")

    add_h2(doc, "（三）项目管理类问题")
    add_body(doc, "问题：HCI 与 VR 报告内容重复（如识别反馈、前情翻页）")
    add_body(doc, "解决方法：制定分工 Excel：VR 写事件与存储，HCI 写文案与 UI；互引 1 段，不展开对方章节。")

    add_body(doc, "问题：AI 生图/全景模块进度与 AR 链路联调时间冲突")
    add_body(doc, "解决方法：先用本地占位图打通保存与阅读链路；接口稳定后再接入真实生图与全景文件。")

    # ═══ 六、总结与展望 ═══
    add_h1(doc, "六、总结与展望")

    add_h2(doc, "（一）工作总结")
    add_body(doc, f"本次 StoryBricks 虚拟现实实验，本人（{STUDENT_NAME}，{STUDENT_ID}）完成了 VR 侧策划与设计全过程，"
             "并主导 F-V03 AR 贴纸叠加、角色识别反馈、故事保存与前情资源管线的开发与测试，"
             "同时承担小组项目安排与 HCI/VR 分工协调。", indent=1)
    add_body(doc, "通过项目实践，我理解了虚实融合不仅是「看见特效」，更是检测稳定性、坐标映射、"
             "事件语义与数据持久化组成的系统工程。ArUco 方案以低硬件成本实现了儿童可感知的「摆对了」反馈；"
             "CompletedStoryStore 则把一次创作会话转化为可回看的完整绘本数据。", indent=1)
    add_body(doc, "测试表明在标记完整可见条件下识别率与贴纸跟随达到课程预期。"
             "创作页按钮布局、语音助手与阅读录音等体验设计详见人机交互课程实验报告；"
             "AI 生图与 360° 全景播放器由组内其他同学完成，本人在接口与保存格式上完成联调。", indent=1)

    add_h2(doc, "（二）未来展望")
    add_numbered(doc, 1, "AR 优化：贴纸平滑与遮挡恢复；多 marker 同时检测性能优化")
    add_numbered(doc, 2, "识别反馈：结合 StoryCreationGapAnalyzer 自动站位描述，减少对手工摆放的依赖")
    add_numbered(doc, 3, "保存扩展：一键导出分享包；云端备份可选方案")
    add_numbered(doc, 4, "全景联调：与组内 360° 模块完善 page_XX_panorama 命名与阅读页加载 fallback")
    add_numbered(doc, 5, "测试深化：补充不同光照/距离下的识别率曲线与长时间运行稳定性")

    add_h1(doc, "七、致谢")
    add_body(doc, "感谢虚拟现实课程老师的悉心指导，为虚实融合方案与测试方法提供了方向指引；"
             "感谢人机交互课程组同学在语音交互、UI 与阅读流程上的协作支持；"
             "感谢组内负责 AI 生图与 360° 全景的同学在接口联调中的配合；"
             "感谢参与测试的同学与家长；感谢 StoryBricks 项目组全体成员的密切配合，使项目顺利落地。")

    doc.save(OUT)
    doc.save(OUT_SUBMIT)
    print(f"Saved: {OUT}")
    print(f"Saved: {OUT_SUBMIT}")


if __name__ == "__main__":
    build()
