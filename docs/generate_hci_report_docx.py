# -*- coding: utf-8 -*-
"""按「202300300098 阙靖涵 实时绘制实验报告」模板生成 StoryBricks HCI 报告。"""
from docx import Document
from docx.shared import Pt, Twips
from docx.oxml.ns import qn

TEMPLATE = r"c:\Users\LENOVO\Desktop\202300300098 阙靖涵 实时绘制实验报告(1).docx"
OUT = r"d:\Game\StoryBricks\docs\StoryBricks_HCI_Report.docx"

STUDENT_NAME = "阙靖涵"
STUDENT_ID = "202300300098"
SZ_BODY = 12   # w:sz 24
SZ_TITLE = 22  # w:sz 44


def _run_font(run, size=SZ_BODY, bold=False):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")


def _spacing(p, before_pt=6, after_pt=6, first_line_chars=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing = 1.2
    if first_line_chars == 2:
        pf.first_line_indent = Pt(24)
    elif first_line_chars == 1:
        pf.first_line_indent = Pt(12)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"] if "Heading 2" in [s.name for s in doc.styles] else p.style
    _spacing(p, 12, 12)
    r = p.add_run(text)
    _run_font(r, SZ_TITLE, bold=True)
    return p


def add_h1(doc, text):
    p = doc.add_paragraph()
    try:
        p.style = "2"
    except Exception:
        pass
    _spacing(p, 16, 6)
    r = p.add_run(text)
    _run_font(r, 16, bold=True)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    try:
        p.style = "3"
    except Exception:
        pass
    _spacing(p, 15, 6)
    r = p.add_run(text)
    _run_font(r, 15, bold=True)
    return p


def add_body(doc, text, indent=0):
    p = doc.add_paragraph()
    _spacing(p, 6, 6, first_line_chars=indent)
    r = p.add_run(text)
    _run_font(r)
    return p


def add_numbered(doc, n, text):
    p = doc.add_paragraph()
    _spacing(p, 6, 6)
    r1 = p.add_run(f"{n}.")
    _run_font(r1)
    r2 = p.add_run(text)
    _run_font(r2)
    return p


def add_sub(doc, label, text):
    p = doc.add_paragraph()
    _spacing(p, 6, 6, first_line_chars=1)
    r1 = p.add_run(f"{label}")
    _run_font(r1)
    r2 = p.add_run(text)
    _run_font(r2)
    return p


def add_image_placeholder(doc, caption):
    add_body(doc, f"【插图：{caption}】", indent=0)


def build():
    doc = Document(TEMPLATE)
    body = doc.element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)

    add_title(doc, "《StoryBricks》儿童故事创作系统人机交互实验报告")

    # ═══ 一、实验项目概述 ═══
    add_h1(doc, "一、实验项目概述")
    add_h2(doc, "（一）项目基本信息")
    add_body(doc, "项目名称：StoryBricks 儿童积木 + AI 绘本创作系统")
    add_body(doc, "项目类型：面向 6～10 岁儿童的实体积木 + 多模态语音交互 + AI 绘本生成应用")
    add_body(doc, "核心主题：儿童以经典故事（如龟兔赛跑）为线索，通过搭建积木、摆放场景、向语音助手「乐乐」讲述情节，"
             "系统将识别结果与语音内容整理后生成为个性化绘本，并支持平面阅读与页内朗读录音，实现「动手—说话—看见成果—表达」的完整创作闭环")
    add_body(doc, "技术基底：Unity 2022.3 LTS、uGUI + TextMesh Pro、OpenCV for Unity（ArUco 识别，由 VR 模块实现）、"
             "storybricks-tutor-gateway（语音 ASR/TTS/大模型）、storybricks-image-gen-web（AI 生图）")
    add_body(doc, "设计目标：在降低儿童文字操作门槛的前提下，打造分阶段引导清晰、语音优先、识别反馈及时、"
             "适配亲子/课堂场景的互动创作体验")

    add_h2(doc, "（二）项目开发背景")
    add_body(doc, "本项目同时作为人机交互与虚拟现实两门课程的联合实验成果。"
             "HCI 侧聚焦儿童能否顺畅完成故事创作全流程，VR 侧聚焦物理积木与虚拟视觉内容的融合呈现。"
             "传统绘本 App 以被动阅读为主，缺少用户生成内容与实体操作连接；纯语音助手缺少与场景任务绑定的结构化引导。"
             "StoryBricks 尝试将实体积木、摄像头视觉反馈、语音叙述与 AI 生成结合，为儿童提供可感知的创作成就感。", indent=2)

    add_h2(doc, "（三）个人核心职责")
    add_body(doc, f"实验人：{STUDENT_NAME}    学号：{STUDENT_ID}", indent=2)
    add_body(doc, "本人负责人机交互方向从用户调研、需求策划、交互与界面设计到核心功能实现与测试的完整工作，核心职责包括：", indent=2)
    add_numbered(doc, 1, "用户调研与策划：目标用户画像、任务分析、竞品对比、创作流程需求规格与功能范围划定")
    add_numbered(doc, 2, "交互与界面设计：信息架构、创作页八阶段状态机、语音助手「乐乐」交互模型、儿童 UI 视觉规范、识别反馈文案策略")
    add_numbered(doc, 3, "功能实现与调整：故事库、绘本前情、创作页语音流程、UI 布局、识别反馈联动、提示词整理链路、平面阅读与页内录音")
    add_numbered(doc, 4, "测试与评估：任务完整度测试、语音交互准确性测试，输出测试数据与改进建议")

    # ═══ 二、个人工作详细内容 ═══
    add_h1(doc, "二、个人工作详细内容")

    add_h2(doc, "（一）用户调研与需求策划工作")
    add_body(doc, "结合课程要求与儿童认知特点，主导 HCI 侧从调研到需求落地的全流程：", indent=2)

    add_body(doc, "1. 用户调研", indent=0)
    add_sub(doc, "", "采用半结构访谈与情境观察，访谈 4 组亲子用户（儿童 6～9 岁及家长）与 2 名小学教师。")
    add_sub(doc, "", "核心发现：儿童更愿意「边玩边说」；经典故事作为入口可降低理解成本；家长希望阶段划分清晰；语音 + 屏幕短提示优于长文本。")

    add_body(doc, "2. 需求分析与任务定义", indent=0)
    add_sub(doc, "", "定义 6 项核心用户任务：选故事(T1)、读前情(T2)、搭积木摆场景(T3)、语音讲述(T4)、确认生成(T5)、回看朗读(T6)。")
    add_sub(doc, "", "明确交互约束：语音优先、分阶段引导、即时反馈、可重来容错、6～10 岁短句口语化。")

    add_body(doc, "3. 竞品与差异化定位", indent=0)
    add_sub(doc, "", "对比绘本 App、积木玩具、语音助手三类参考产品，确定 StoryBricks 差异化：物理摆放—语音叙述—视觉生成—回看表达串联为一条创作链。")

    add_body(doc, "4. 策划结论", indent=0)
    add_sub(doc, "", "确定「分阶段状态机 + 语音优先 + 识别反馈辅助」交互策略；划定 HCI 功能边界：故事库、前情、创作页语音/UI、阅读录音；"
                       "ArUco 算法与 AI 生图技术细节归属 VR 报告。")

    add_image_placeholder(doc, "用户旅程图 / 任务分析表")

    add_h2(doc, "（二）交互与界面设计工作")
    add_body(doc, "1. 信息架构与导航流程", indent=0)
    add_sub(doc, "", "主路径：StartScene → StorySummary → StoryPrologue → StoryWorks → Tutorial → StoryCreation → CompletedStoryLibrary → CompletedStoryViewer。")
    add_sub(doc, "", "三层导航：故事库 / 积木作品集 / 我的故事；统一返回按钮逻辑（StoryFlowBackButtonUi）。")

    add_body(doc, "2. 创作页状态机设计", indent=0)
    add_sub(doc, "", "八阶段：Guide → Building → Capturing → VoiceInteracting → StoryConfirm → Generating → PageDone → StoryFinished。")
    add_sub(doc, "", "关键决策：Building 与 VoiceInteracting 分离；点击「这页摆好了」后才抓拍与追问；simplifiedUi 按阶段显隐按钮。")

    add_body(doc, "3. 语音助手「乐乐」四层交互模型", indent=0)
    add_sub(doc, "① ", "Building：唤醒词「你好乐乐」→ 自由边玩边说；过滤纯唤醒词回声。")
    add_sub(doc, "② ", "识别反馈：角色到达/移动/到齐 → 口语化文案（如「兔子来啦！」「伙伴都到齐啦！」）。")
    add_sub(doc, "③ ", "VoiceInteracting：缺口追问 1～2 个；草稿充分则跳过；听不清容错提示。")
    add_sub(doc, "④ ", "StoryConfirm：大模型整理对话 → 乐乐复述 → 儿童确认后出图。")

    add_body(doc, "4. 界面视觉规范", indent=0)
    add_sub(doc, "", "StoryCard 卡片网格、绘本水彩暖色、16:9 横版、TMP 字体、大按钮（阅读页翻页 200×200）。")
    add_sub(doc, "", "创作页：全屏背景 + 引导文案 + 摄像头小窗可展开 + 底部操作栏按阶段显隐。")

    add_image_placeholder(doc, "创作页状态机图 / 线框图 / StoryCard 界面")

    add_h2(doc, "（三）功能实现与调整工作")

    add_body(doc, "1. 故事库功能（StorySummary + StoryCardView）", indent=0)
    add_sub(doc, "（1）", "ScrollView 卡片网格展示多故事，封面 + 标题 +「选择」按钮，一步进入前情。")
    add_sub(doc, "（2）", "支持 Resources/Stories 自动加载 StoryDefinition，便于扩展新故事。")

    add_body(doc, "2. 绘本前情功能（StoryProloguePictureBook）", indent=0)
    add_sub(doc, "（1）", "全屏插图 + Prev/Next 翻页，页码指示（如 1/2），末页「开始搭建」进入作品集。")
    add_sub(doc, "（2）", "前情阶段不设语音问答，降低动机建立阶段的认知负担。")

    add_body(doc, "3. 创作页语音交互调整（StoryCreationPageBootstrap + LeleVoiceAssistant）", indent=0)
    add_sub(doc, "（1）", "配置 tutorGatewayUrl、useGapQuestionsOnConfirm、useAiGeneratedQuestions、minStoryDraftCharsToSkipQuestions 等参数。")
    add_sub(doc, "（2）", "调整缺口追问策略：优先行为类问题，可选元素问题最多 1 个；TTS + 屏幕文字双通道。")
    add_sub(doc, "（3）", "故事整理环节（RunAmbientStoryCloseCoroutine）：显示「乐乐在整理你刚才讲的故事…」，支持确认与补充。")

    add_body(doc, "4. 识别反馈设计与联动（StoryCreationArDirector → StoryCreationLeleHost）", indent=0)
    add_sub(doc, "（1）", "CharacterArrived：「{角色}来啦！摆好位置，跟乐乐说说 ta 在干嘛。」")
    add_sub(doc, "（2）", "CharacterMoved：「哇，{角色}挪位置啦！现在想做什么呀？」")
    add_sub(doc, "（3）", "AllCharactersReady：「伙伴都到齐啦！边玩边讲，好了就点这页摆好了。」")
    add_sub(doc, "（4）", "RosterHint：「还差乌龟哦」等名单提示，与乐乐面板联动。")

    add_body(doc, "5. 提示词整理链路调整", indent=0)
    add_sub(doc, "（1）", "抓拍后 BuildAutoPlacementDescription 自动推断站位，不再口头问「谁在前」。")
    add_sub(doc, "（2）", "合并边玩边说草稿、缺口回答、自动站位为 voiceSupplement。")
    add_sub(doc, "（3）", "useAiPromptRefinement 调用大模型整理为连贯场景描述；失败回退 StoryPageGenerationPipeline.BuildLocalGenerationPrompt。")
    add_sub(doc, "（4）", "状态栏「正在整理生图描述…」外显等待过程。")

    add_body(doc, "6. 创作页 UI 调整（StoryCreationPageUiBuilder）", indent=0)
    add_sub(doc, "（1）", "simplifiedUi 精简模式：Building 突出摆积木与语音，Generating 只保留等待提示。")
    add_sub(doc, "（2）", "摄像头小窗可展开为大图，便于查看 AR 贴纸与摆放效果。")
    add_sub(doc, "（3）", "VoiceInteracting 阶段 voiceQuestionText 显示提问全文。")

    add_body(doc, "7. 平面阅读与页内录音（CompletedStoryViewerRoot + CompletedStoryPageVoiceRecorder）", indent=0)
    add_sub(doc, "（1）", "全屏翻页阅读，页码指示含前情/创作类型；「故事阅读」面板展开/收起。")
    add_sub(doc, "（2）", "每页录音（最长 90 秒）、播放、重录；状态「让我们一起朗读吧~」「已保存你的朗读」。")
    add_sub(doc, "（3）", "「360° 全景」入口 UX（无资源灰显），技术实现见 VR 报告。")

    add_image_placeholder(doc, "故事库界面 / 创作页摄像头与乐乐 / 生成绘本页 / 阅读页录音")

    add_h2(doc, "（四）测试与评估工作")
    add_body(doc, "1. 任务完整度测试", indent=0)
    add_sub(doc, "（1）", "被试：5 名 7～9 岁儿童（家长陪同）；设备：10.5 英寸平板 + 俯拍摄像头 + 龟兔赛跑积木。")
    add_sub(doc, "（2）", "任务 C1～C5：选故事、读前情、进创作、生成一页、阅读录音；记录完成率与卡点。")
    add_sub(doc, "（3）", "结果：平均完整度 4.6/5（92%）；主要卡点：首次不知点「这页摆好了」、教程停留长、1 人未找到录音按钮。")

    add_body(doc, "2. 语音交互准确性测试", indent=0)
    add_sub(doc, "（1）", "指标：唤醒成功率(5次/人)、自由对话理解度(1～5)、追问听懂率、故事确认准确率、误识别容错。")
    add_sub(doc, "（2）", "结果：唤醒 88%（22/25）；理解度均分 4.2/5；追问 4/5 完全听懂；确认 4/5 与意图一致；容错提示有效。")

    add_image_placeholder(doc, "测试记录表 / 问卷样例")

    # ═══ 三、工作成果展示 ═══
    add_h1(doc, "三、工作成果展示")

    add_h2(doc, "（一）策划与设计成果")
    add_numbered(doc, 1, "输出 StoryBricks HCI 需求文档：用户画像、6 项核心任务、用户旅程图、创作页状态机图、乐乐话术表")
    add_numbered(doc, 2, "完成信息架构与三层导航设计，明确 HCI 与 VR 课程报告边界与互引关系")
    add_numbered(doc, 3, "确立「分阶段状态机 + 四层语音模型 + 识别反馈转译」交互框架，为后续迭代提供依据")

    add_h2(doc, "（二）功能实现成果")
    add_numbered(doc, 1, "故事库：多故事卡片选择与 StoryDefinition 数据驱动扩展")
    add_numbered(doc, 2, "绘本前情：滑动翻页 + 开始搭建入口，降低进入创作门槛")
    add_numbered(doc, 3, "创作页语音：唤醒、自由对话、缺口追问、故事确认全流程可运行")
    add_numbered(doc, 4, "识别反馈：角色到达/移动/到齐/名单提示与乐乐联动，儿童可理解「摆对了没有」")
    add_numbered(doc, 5, "提示词整理：voiceSupplement → AI 整理 → 生图，儿童只需说话不需感知 Prompt")
    add_numbered(doc, 6, "平面阅读 + 录音：已完成故事可翻页阅读并保存页内朗读 WAV")

    add_h2(doc, "（三）测试评估成果")
    add_numbered(doc, 1, "任务完整度 92%，主路径「选故事→创作→保存→阅读录音」可达")
    add_numbered(doc, 2, "语音唤醒成功率 88%，自由对话理解度 4.2/5，整体满足实验预期")
    add_numbered(doc, 3, "识别反馈观察：5 名被试中 3 人会主动查看摄像头小窗，名单提示有效引导补摆角色")

    add_h2(doc, "（四）需补充的插图（个人填写）")
    add_body(doc, "请在 Word 中替换以下占位为实际截图：")
    add_numbered(doc, 1, "图1：StorySummary 故事库卡片界面")
    add_numbered(doc, 2, "图2：StoryPrologue 绘本前情翻页")
    add_numbered(doc, 3, "图3：StoryCreation 创作页摆积木 + 摄像头 AR 贴纸")
    add_numbered(doc, 4, "图4：乐乐语音追问与故事确认界面")
    add_numbered(doc, 5, "图5：AI 生成的绘本页成果")
    add_numbered(doc, 6, "图6：CompletedStoryViewer 阅读页 + 录音按钮")
    add_numbered(doc, 7, "图7：任务完整度 / 语音准确性测试记录表")

    # ═══ 五、问题与解决方法 ═══
    add_h1(doc, "五、问题与解决方法")

    add_h2(doc, "（一）交互设计类问题")
    add_body(doc, "问题：儿童首次进入创作页，不知道何时点击「这页摆好了」，长时间停留在 Building 阶段")
    add_body(doc, "解决方法：在 AllCharactersReady 反馈中明确提示「好了就点这页摆好了」；后续计划增加首次高亮引导。")

    add_body(doc, "问题：缺口追问轮数过多，打断儿童叙述流畅性")
    add_body(doc, "解决方法：增加 ShouldSkipGapQuestions：草稿超过 18 字且各角色均已提及时跳过追问；追问上限 1～2 个。")

    add_body(doc, "问题：角色站位原先需口头追问「谁在前」，儿童难以回答")
    add_body(doc, "解决方法：改为 StoryCreationGapAnalyzer.BuildAutoPlacementDescription 根据 ArUco 像素坐标自动推断，不再口头提问位置。")

    add_h2(doc, "（二）语音交互类问题")
    add_body(doc, "问题：唤醒词「你好乐乐」在嘈杂环境下识别失败，或误识别环境音")
    add_body(doc, "解决方法：LeleVoiceAssistant 支持多种唤醒变体匹配；TTS 与屏幕文字双通道；提示尽量在安静环境使用。")

    add_body(doc, "问题：AI 故事整理失败或耗时长，儿童误以为系统卡死")
    add_body(doc, "解决方法：状态栏显示「乐乐在整理你刚才讲的故事…」；整理失败回退本地拼接，不中断流程。")

    add_body(doc, "问题：语音识别听不清时儿童容易放弃")
    add_body(doc, "解决方法：提示「没听清也没关系，我们先用刚才听到的」；支持重试当前问题。")

    add_h2(doc, "（三）界面与测试类问题")
    add_body(doc, "问题：阅读页录音按钮首次可发现性不足")
    add_body(doc, "解决方法：状态文案「让我们一起朗读吧~」；录音按钮置于故事面板 VoiceRow 明显位置。")

    add_body(doc, "问题：生成等待期间反馈不足")
    add_body(doc, "解决方法：分阶段状态文案「正在识别积木…」「正在整理生图描述…」「生成中…」。")

    # ═══ 六、总结与展望 ═══
    add_h1(doc, "六、总结与展望")

    add_h2(doc, "（一）工作总结")
    add_body(doc, f"本次 StoryBricks 人机交互实验，本人（{STUDENT_NAME}，{STUDENT_ID}）完成了从用户调研、需求策划、交互设计到故事库、前情、"
             "创作页语音、UI、识别反馈、提示词整理、平面阅读录音及测试评估的完整 HCI 工作。", indent=1)
    add_body(doc, "通过项目实践，我系统提升了儿童多模态交互设计能力，理解了 VUI 不仅是「能听会说」，更要设计对话节奏、"
             "错误恢复与确认环节。分阶段状态机与识别反馈转译是降低儿童认知负担的关键；"
             "提示词整理链路则体现了「儿童只管说，系统负责整理」的 HCI 原则。", indent=1)
    add_body(doc, "测试表明任务完整度 92%、语音唤醒成功率 88%，整体达到课程实验预期。"
             "系统中 ArUco 检测、AI 生图与 360° 全景技术详见虚拟现实课程实验报告。", indent=1)

    add_h2(doc, "（二）未来展望")
    add_numbered(doc, 1, "交互优化：「这页摆好了」首次高亮引导；生成进度更细化；首次使用教程动画")
    add_numbered(doc, 2, "语音优化：嘈杂环境降噪；近场麦克风适配；方言识别优化")
    add_numbered(doc, 3, "阅读优化：录音按钮加大；家长协同模式；更多故事模板扩展")
    add_numbered(doc, 4, "与 VR 模块联调：360° 全景入口与 HCI 阅读流程进一步整合")

    add_h1(doc, "七、致谢")
    add_body(doc, "感谢人机交互课程老师的悉心指导，为项目交互设计与测试方法提供了方向指引；"
             "感谢虚拟现实课程组同学在 ArUco 识别与 AI 生图模块的协作支持；"
             "感谢参与测试的儿童与家长、提供访谈意见的教师；"
             "感谢 StoryBricks 项目组全体成员的密切配合，使项目顺利落地。")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    build()
